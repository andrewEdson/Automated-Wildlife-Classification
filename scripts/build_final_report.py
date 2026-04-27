"""Generate Final Report.docx in the project folder.

Reuses the midterm report's text where it remains accurate; adds new
Methodology, Results, Discussion, and Future Work sections that match
what was actually implemented.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUT  = ROOT / "Final Report.docx"

V0  = ROOT / "results" / "species" / "baseline_v0"
V1  = ROOT / "results" / "species_cropped"

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H(text, level=1):
    """Bold heading like the midterm — no Word styles, just bold + larger."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14 if level == 1 else 12)


def P(text):
    doc.add_paragraph(text)


def B(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def IMG(path, width_in=5.5, caption=None):
    if not path.exists():
        P(f"[missing figure: {path.name}]"); return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)


# ── TITLE / TEAM ────────────────────────────────────────────────────────────
H("Title:")
P("Automated Wildlife Classification Using Convolutional Neural Networks on Camera "
  "Trap Images")

H("Team Information:")
B(["Andrew Edson – Student ID: 11778881"])

# ── ABSTRACT ────────────────────────────────────────────────────────────────
H("Abstract:")
P("Monitoring wildlife populations is essential for conservation and environmental "
  "research. Manual analysis of camera trap images is labor intensive and time consuming. "
  "This project builds an automated classification system using a Convolutional Neural "
  "Network to identify and classify animal species. The approach leverages transfer "
  "learning on ResNet-50 pretrained on ImageNet, fine-tuned on the Caltech Camera Traps "
  "(CCT20) benchmark, which contains around 243,100 images from 140 camera locations in "
  "the Southwestern United States. Two versions of the model were trained and compared. "
  "The baseline used the full camera-trap frame as input with a frozen backbone and "
  "achieved a macro F1 score of 0.219 on the out-of-domain trans-test split. The improved "
  "version cropped each image to its ground-truth animal bounding box at 384 px, "
  "fine-tuned the last residual block, and used class-balanced cross-entropy loss. This "
  "version achieved a macro F1 score of 0.388 on the same trans-test split, a relative "
  "improvement of about 77 percent. Results show that the dominant failure mode of the "
  "baseline was the model learning camera-location backgrounds rather than the animal "
  "itself, and that simply giving the classifier a tighter view of the subject is more "
  "valuable than further architectural changes at this scale.")

# ── INTRODUCTION ────────────────────────────────────────────────────────────
H("Introduction:")
P("Wildlife monitoring provides important insight for conservation, wildlife management, "
  "and even construction planning. Camera traps are widely used across the United States "
  "to capture images of animals in their natural habitats. However, the data collected "
  "must usually be labeled manually, which requires a large amount of human time and "
  "labor.")
P("Automating species classification using deep learning can greatly reduce this "
  "workload. Convolutional Neural Networks (CNNs) have been shown to perform very well "
  "on image classification tasks because they are able to learn hierarchical visual "
  "features such as edges, textures, shapes, and patterns.")
P("In this project, transfer learning is used with a pretrained model, ResNet-50. This "
  "allows the network to reuse the general visual features it has already learned and "
  "adapt them to the camera trap dataset.")
P("However, this dataset presents several challenges. The Caltech Camera Traps dataset "
  "is highly imbalanced, as most of the images are empty and contain no animals. Even "
  "among the animal images there is a large discrepancy in frequency, animals such as "
  "opossums and rabbits appear very often, while others such as fox and badger appear "
  "only a handful of times. This makes accurate classification of less common species "
  "much more difficult than classifying frequently seen animals. A second challenge is "
  "that camera-trap animals usually occupy only a small portion of a wide-angle frame, "
  "so a 224-pixel center crop of the full image often loses the subject. Addressing this "
  "by cropping the image to the animal before classification turned out to be the single "
  "most important change made in this project.")

# ── LITERATURE REVIEW ──────────────────────────────────────────────────────
H("Literature Review:")
P("Camera traps are commonly used in wildlife monitoring but generate large numbers of "
  "images that must be manually labeled, many of which contain no animals. This occurs "
  "because the cameras are triggered by movement and heat, causing environmental factors "
  "to produce false activations. Automating this process has become an active area of "
  "research in computer vision. Recent studies show that Convolutional Neural Networks "
  "(CNNs) can accurately classify species in camera-trap images despite challenges such "
  "as poor lighting, occlusion, and complex backgrounds (Beery et al., 2018; Schneider "
  "et al., 2018).")
P("Because camera-trap datasets are smaller than standard vision datasets, transfer "
  "learning is often applied. A network pre-trained on a large dataset such as ImageNet "
  "can be fine-tuned to recognize wildlife species by leveraging previously learned "
  "visual features. Residual Networks (ResNet) are commonly used for this purpose due to "
  "their ability to train deep architectures effectively (He et al., 2015).")
P("A recurring observation in the camera-trap literature is that classifiers trained "
  "directly on full frames tend to overfit to the specific background of each camera "
  "location. Beery et al. (2018) showed that performance on new locations (the "
  "\"trans\" split) can drop sharply compared to performance on familiar locations. "
  "A common way to address this is to first detect the animal with a generic detector "
  "and only classify the cropped subject. Beery et al. (2019) introduced an efficient "
  "pipeline based on a dedicated camera-trap detector (MegaDetector) and showed that "
  "this dramatically improves species classification.")
P("Deep learning approaches have also demonstrated strong performance in wildlife "
  "classification tasks. Norouzzadeh et al. (2018) showed that convolutional neural "
  "networks can automatically identify and count animal species in camera-trap images "
  "with accuracy comparable to humans. Automating this process represents a significant "
  "improvement in efficiency, reducing the time required to manually label large "
  "datasets and allowing ecologists and conservationists to analyze wildlife populations "
  "at much larger scales.")
P("Another challenge in camera-trap data is severe class imbalance. Standard inverse-"
  "frequency loss weighting can place very large weights on classes with only a few "
  "samples, which adds noise to training rather than improving recognition. Cui et al. "
  "(2019) proposed a class-balanced loss based on the effective number of samples that "
  "is much gentler on the long tail; this loss is used in the improved version of the "
  "model trained for this project.")
P("This project uses the Caltech Camera Traps dataset (CCT20 subset), which contains "
  "labeled wildlife and empty images captured by camera traps along with ground-truth "
  "animal bounding boxes for most annotations, making it suitable for evaluating "
  "CNN-based animal classification under realistic location shift.")

# ── METHODOLOGY ────────────────────────────────────────────────────────────
H("Methodology:")
P("The original midterm proposal described a two-stage pipeline: a binary "
  "empty-vs-animal filter followed by a species classifier. After working with the "
  "CCT20 subset directly, the empty filter turned out to be unnecessary for the "
  "species training set, because the official CCT20 train split contains no empty "
  "frames. Empty frames only appear in the validation and test splits, where they "
  "can be filtered out using the annotations. Because of this, the project was "
  "simplified to a single multi-class species classifier and the engineering effort "
  "was redirected to two more impactful directions: cropping the input to the animal "
  "and properly fine-tuning the backbone.")
H("1. Dataset", level=2)
P("The dataset used is the CCT20 subset of Caltech Camera Traps, which is split into "
  "in-domain (cis) and out-of-domain (trans) evaluation. The trans split is taken from "
  "camera locations the model has never seen during training, so it is the more "
  "honest measure of generalization. Annotations are stored in COCO-style JSON files "
  "and most of them include ground-truth bounding boxes for the animal in the image "
  "(about 90% of the train split, 81% of cis-test, and 78% of trans-test).")
H("2. Baseline Classifier (Version 0)", level=2)
P("The baseline classifier is a ResNet-50 with weights pretrained on ImageNet. The "
  "convolutional backbone is frozen and only the final fully connected layer is "
  "replaced with a 15-class head. Inputs are full camera-trap frames resized to "
  "224 pixels. Training uses Adam at a learning rate of 1e-3, cross-entropy loss with "
  "inverse-frequency class weights, and a WeightedRandomSampler to keep mini-batches "
  "balanced across species. This is the model that is described in the original "
  "midterm proposal, with one change: rather than two stages, all 15 species "
  "categories are predicted in a single model.")
H("3. Improved Classifier (Version 1: GT-Crop, 384 px, Fine-tuned)", level=2)
P("Version 1 makes four changes relative to the baseline, each motivated by an "
  "observed failure mode of the baseline:")
B([
    "Input is the ground-truth animal bounding box, padded by 10% on each side and "
    "resized to 384 pixels. This addresses the dominant failure of the baseline, which "
    "was learning the background of the camera location instead of the animal.",
    "The last residual block (layer4) of the ResNet-50 backbone is fine-tuned at a "
    "smaller learning rate (1e-4) while the head is trained at the usual 1e-3 "
    "(discriminative learning rate). Earlier layers remain frozen.",
    "Cross-entropy loss uses class-balanced weighting (Cui et al., 2019) with beta = "
    "0.999, instead of raw inverse-frequency. This is gentler on classes with very few "
    "samples and reduces noise in the gradient.",
    "Three species with too few training samples to learn (badger n=9, deer n=45, "
    "fox n=5) are dropped from training and evaluation. The headline metric becomes "
    "12-class macro F1.",
])
P("Animal images that do not have a ground-truth bounding box are also dropped (about "
  "10% of train). After all filtering the splits used for Version 1 contain 12,559 "
  "training crops, 1,760 validation crops, 13,118 cis-test crops, and 18,770 "
  "trans-test crops.")
H("4. Training", level=2)
P("Training uses a batch size of 32, runs for up to 25 epochs with early stopping "
  "patience of 7 on validation macro F1, and reduces the learning rate by 0.3x after "
  "3 epochs without improvement. The first 2 epochs run with layer4 frozen so that "
  "the randomly initialized head can settle before backprop reaches the backbone, "
  "after which layer4 unfreezes for fine-tuning. Training was performed on a Mac "
  "with an Apple M-series GPU using the PyTorch MPS backend.")
H("5. Evaluation", level=2)
P("Both versions are evaluated on accuracy, precision, recall, F1, and confusion "
  "matrices for each of the cis-test and trans-test splits. Macro-averaged F1 is "
  "used as the headline metric because it weights all species equally regardless of "
  "their frequency, and is therefore a more honest measure on a heavily imbalanced "
  "dataset than overall accuracy.")

# ── RESULTS ────────────────────────────────────────────────────────────────
H("Results:")
P("The baseline classifier (Version 0, frozen backbone, full-frame 224 px input, "
  "15 classes) achieved an out-of-domain trans-test macro F1 of 0.219. The per-class "
  "F1 breakdown showed that the model performed reasonably on a handful of large or "
  "visually distinct classes (raccoon, squirrel, car) but scored near zero on most "
  "species, including several with thousands of training samples. The confusion "
  "matrix in Figure 1 makes the failure mode clear: predictions collapsed onto a few "
  "common classes, with very heavy off-diagonal mass. This is the behavior expected "
  "when the model latches onto camera-location backgrounds rather than the animal "
  "itself.")

IMG(V0 / "confusion_matrix.png", width_in=5.5,
    caption="Figure 1. Baseline (v0) confusion matrix on the out-of-domain trans-test "
            "split. Heavy off-diagonal mass indicates predictions collapsing onto a "
            "few common classes.")
IMG(V0 / "per_class_f1.png", width_in=6.0,
    caption="Figure 2. Baseline (v0) per-class F1 on trans-test. Macro F1 = 0.219.")

P("Version 1 (GT-cropped, 384 px, layer4 fine-tuned, 12 classes, class-balanced loss) "
  "raised trans-test macro F1 to 0.388. On the in-domain cis-test split, macro F1 "
  "reached 0.631. The training curves are shown in Figure 3 and the per-class F1 "
  "breakdown is shown in Figure 4.")

IMG(V1 / "training_curves.png", width_in=6.0,
    caption="Figure 3. Version 1 training curves. Validation macro F1 climbs to 0.80 "
            "on the in-domain validation split.")
IMG(V1 / "per_class_f1.png", width_in=6.0,
    caption="Figure 4. Version 1 per-class F1 on trans-test. Macro F1 = 0.388. Car, "
            "opossum, raccoon, and dog are now reliably classified, while bird, cat, "
            "and rodent remain difficult.")
IMG(V1 / "confusion_matrices.png", width_in=6.5,
    caption="Figure 5. Version 1 confusion matrices for cis-test (left) and trans-test "
            "(right). Most predictions are now on the diagonal, but a residual "
            "domain-shift gap remains between cis and trans.")

P("A side-by-side comparison of the two versions on the same trans-test split is "
  "given in Table 1. Because Version 1 is evaluated on 12 classes (with three "
  "very-low-data classes dropped), a fair version of the baseline number recomputed "
  "over the same 12 classes is also reported.")

# Comparison table
tbl = doc.add_table(rows=4, cols=4)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Model"
hdr[1].text = "Cis-test macro F1"
hdr[2].text = "Trans-test macro F1"
hdr[3].text = "Classes"
rows = [
    ("v0 (full frame, frozen, raw inverse-freq CE)", "—",     "0.219", "15"),
    ("v0 reweighted to same 12 classes",              "—",     "≈ 0.274", "12"),
    ("v1 (GT-crop 384 px, layer4 fine-tuned, CB loss)", "0.631", "0.388", "12"),
]
for i, r in enumerate(rows, start=1):
    cells = tbl.rows[i].cells
    for j, val in enumerate(r):
        cells[j].text = val

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
crun = cap.add_run("Table 1. Headline comparison of baseline and improved classifiers.")
crun.italic = True; crun.font.size = Pt(10)

P("Sample crops fed into Version 1 are shown in Figure 6. Because the input is the "
  "padded ground-truth bounding box, the animal occupies most of the frame and the "
  "model can no longer rely on background context.")

IMG(V1 / "sample_grid_cropped.png", width_in=5.5,
    caption="Figure 6. Sample crops (padded ground-truth bounding boxes) used as "
            "input to the Version 1 classifier.")

# ── DISCUSSION ─────────────────────────────────────────────────────────────
H("Discussion:")
P("The largest single gain in this project came not from a new architecture or longer "
  "training, but from changing what the classifier actually looks at. The baseline was "
  "given a 224-pixel center crop of an entire wide-angle camera-trap frame. In most of "
  "those frames the animal is small and off-center, and the camera-location background "
  "(specific rocks, branches, fence posts, trail surfaces) is large, sharp, and "
  "consistent across all images from that location. Under those conditions the model "
  "has a strong incentive to learn \"this is the bobcat camera\" rather than \"this is "
  "a bobcat,\" and that is exactly what the trans-test macro F1 of 0.219 suggests was "
  "happening: as soon as the model was tested on locations it had never seen, accuracy "
  "collapsed.")
P("Version 1 attacks this directly by cropping each image to its ground-truth bounding "
  "box before the classifier ever sees it. This nearly doubled trans-test macro F1, "
  "from 0.219 to 0.388. At the same time, fine-tuning the last residual block of the "
  "ResNet-50 backbone (rather than only training a linear head on frozen ImageNet "
  "features) and switching to class-balanced cross-entropy contributed the rest of the "
  "improvement. None of these changes are exotic: they are the standard knobs from "
  "the recent camera-trap literature applied carefully.")
P("The cis vs trans gap (0.631 vs 0.388) is still large, which means a meaningful "
  "amount of overfitting to the cis camera locations remains even with cropped inputs. "
  "Some of this is unavoidable on CCT20: the trans split was deliberately constructed "
  "to be hard. Some of it is also tied to specific species. Cars, opossums, raccoons, "
  "and dogs are now classified reliably (F1 above 0.5 on trans-test), but bird (F1 "
  "0.10), cat (0.21), and rodent (0.00) remain weak. Bird and rodent both have very "
  "few trans-test samples (180 and 26 respectively), so their numbers are noisy. Cat, "
  "with 1,366 trans-test samples, is a real failure and is most often confused with "
  "bobcat - a known hard case in the camera-trap literature because the two species "
  "look quite similar in nighttime infrared imagery.")
P("A limitation of the Version 1 evaluation is that it uses ground-truth bounding "
  "boxes at test time. In a real deployment the model would be paired with an animal "
  "detector such as MegaDetector, which would produce slightly noisier crops. The "
  "trans-test number reported here is therefore best read as the upper bound that the "
  "classifier could achieve in a deployed pipeline that uses a strong detector front "
  "end. Plans to actually run that end-to-end pipeline are described in the next "
  "section.")

# ── FUTURE WORK ────────────────────────────────────────────────────────────
H("Future Work:")
P("The most obvious next step is to replace the ResNet-50 backbone with a more modern "
  "one. ResNet-50 at 384 pixels is a strong baseline, but recent backbones such as "
  "ConvNeXt-Tiny or a ViT-S using DINOv2 features would likely give another large "
  "gain on this kind of fine-grained classification, especially at higher input "
  "resolutions where small features (ear shape, tail markings) start to matter.")
P("A second direction is to evaluate the full deployed pipeline using MegaDetector "
  "(Beery et al., 2019) to crop animals at inference time, instead of using "
  "ground-truth boxes. This would give a more honest deployment number and also let "
  "the classifier handle the small fraction of images that do not have a ground-truth "
  "box in the dataset. Initial setup for this was completed during the project: the "
  "MegaDetector v5a model was successfully loaded and run on a 50-image smoke test, "
  "producing tight detections on small subjects that would be invisible in a 224-pixel "
  "center crop. The full inference pass over the dataset was deferred so that "
  "training time could be spent on the classifier itself, but the scripts are in "
  "place for a follow-up run.")
P("Third, CCT camera triggers fire in bursts of three frames at the same location "
  "with the same animal. The current evaluation treats each frame independently. "
  "Averaging softmax probabilities across the burst at test time would likely give a "
  "small but free improvement to macro F1, since one of the three frames is often "
  "much clearer than the others.")
P("Finally, the species that remain hard for Version 1 (bird, cat, rodent) are also "
  "the species with the fewest trans-test samples or the most confusable visual "
  "features. Targeted data augmentation, hierarchical labels (for example, grouping "
  "cat and bobcat under a shared parent class for an auxiliary loss), and pulling in "
  "additional camera-trap data from related datasets are all reasonable directions "
  "beyond the scope of this course.")

# ── REFERENCES ─────────────────────────────────────────────────────────────
H("References:")
P("Beery, S., Van Horn, G., & Perona, P. (2018). Recognition in Terra Incognita. "
  "ECCV.")
P("Beery, S., Morris, D., & Yang, S. (2019). Efficient Pipeline for Camera Trap Image "
  "Review. Proceedings of the IEEE Conference on Computer Vision and Pattern "
  "Recognition (CVPR) Workshops.")
P("Cui, Y., Jia, M., Lin, T.-Y., Song, Y., & Belongie, S. (2019). Class-Balanced Loss "
  "Based on Effective Number of Samples. Proceedings of the IEEE/CVF Conference on "
  "Computer Vision and Pattern Recognition (CVPR).")
P("He, K., Zhang, X., Ren, S., & Sun, J. (2015). Deep residual learning for image "
  "recognition. CVPR.")
P("Norouzzadeh, M. S., Nguyen, A., Kosmala, M., Swanson, A., Palmer, M., Packer, C., "
  "& Clune, J. (2018). Automatically identifying, counting, and describing wild "
  "animals in camera-trap images with deep learning. Proceedings of the National "
  "Academy of Sciences, 115(25), E5716-E5725.")
P("Schneider, S., Taylor, G. W., & Linquist, S. (2018). Deep learning object "
  "detection methods for ecological camera trap data. Ecological Informatics.")
P("Caltech Camera Traps Dataset. https://lila.science/datasets/caltech-camera-traps/")

doc.save(OUT)
print(f"Wrote {OUT}")
